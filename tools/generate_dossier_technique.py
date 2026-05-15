# -*- coding: utf-8 -*-
"""Génère le Dossier technique ServiDom (Word). Exécution : python generate_dossier_technique.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def main():
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Dossier_technique_ServiDom.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Page de garde ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DOSSIER TECHNIQUE\n(Dossier de réalisation)")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ServiDom")
    r.bold = True
    r.font.size = Pt(22)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        "Application mobile et API REST pour la mise en relation\n"
        "entre clients et prestataires de services à domicile."
    )

    doc.add_paragraph()
    for line in [
        "Formation / institution : [À compléter]",
        "Module / séminaire : MIDA – [À compléter]",
        "Groupe : [À compléter]",
        "Membres du projet : [À compléter]",
        "Année académique : 2025–2026",
        "Date de version du document : [À compléter]",
        "Encadrant : [À compléter]",
    ]:
        pi = doc.add_paragraph(line)
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Table des matières (manuelle) ---
    add_heading(doc, "Table des matières", 1)
    toc = [
        "1. Introduction",
        "2. Analyse du besoin",
        "3. Conception",
        "4. Réalisation",
        "5. Déploiement et exploitation",
        "6. Tests et validation",
        "7. Difficultés rencontrées et solutions",
        "8. Conclusion et perspectives",
        "Annexes",
    ]
    for e in toc:
        doc.add_paragraph(e, style="List Number")

    doc.add_page_break()

    # --- 1. Introduction ---
    add_heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "ServiDom est une solution logicielle visant à faciliter l’accès aux services à domicile "
        "(ménage, cuisine, jardinage, plomberie, électricité, sécurité, etc.) en permettant aux "
        "clients de découvrir des prestataires, de consulter leurs offres et de passer des réservations. "
        "Le périmètre actuel couvre une application mobile (Flutter) consommant une API REST (Node.js / Express) "
        "persistant les données dans PostgreSQL.",
    )
    add_para(
        doc,
        "Les objectifs principaux sont : simplifier la recherche de prestataires fiables ; "
        "standardiser le processus de réservation (date, durée, adresse, montant) ; "
        "offrir un suivi des statuts de réservation et un mécanisme d’avis après prestation.",
    )
    add_para(
        doc,
        "Livrables associés au présent dossier : description de l’architecture, du modèle de données, "
        "des principales routes API, des prérequis et procédures d’installation, ainsi qu’une synthèse "
        "des tests et perspectives d’évolution.",
    )

    # --- 2. Analyse du besoin ---
    add_heading(doc, "2. Analyse du besoin", 1)
    add_heading(doc, "2.1 Acteurs", 2)
    add_bullets(
        doc,
        [
            "Client : s’inscrit, consulte les prestataires et services, crée des réservations, consulte l’historique, peut laisser un avis.",
            "Prestataire : s’inscrit, propose des services liés à une catégorie, reçoit des demandes de réservation, met à jour le statut des interventions.",
            "Administrateur : prévu dans le modèle de rôles (valeur « admin » en base) ; les écrans dédiés peuvent être étendus selon les besoins du déploiement.",
        ],
    )
    add_heading(doc, "2.2 Cas d’usage synthétiques", 2)
    add_bullets(
        doc,
        [
            "S’authentifier (inscription / connexion) et conserver une session via jeton JWT.",
            "Mettre à jour son profil (nom, prénom, email, quartier, coordonnées géographiques optionnelles).",
            "Lister les catégories et les prestataires ; afficher le détail d’un prestataire.",
            "Créer un service (prestataire authentifié) avec possibilité d’upload d’image.",
            "Créer une réservation (client), consulter « mes réservations », modifier le statut, déposer un avis.",
        ],
    )
    add_heading(doc, "2.3 Contraintes", 2)
    add_bullets(
        doc,
        [
            "Connexion réseau entre le mobile et le serveur API (CORS activé côté serveur).",
            "Sécurité : mots de passe hashés (bcrypt), routes protégées par middleware JWT.",
            "Cohérence des données : clés étrangères PostgreSQL entre utilisateurs, services, réservations et avis.",
        ],
    )

    # --- 3. Conception ---
    add_heading(doc, "3. Conception", 1)
    add_heading(doc, "3.1 Architecture logicielle", 2)
    add_para(
        doc,
        "L’architecture retenue est une architecture client–serveur classique : "
        "l’application Flutter joue le rôle de client lourd ; le backend expose une API REST JSON ; "
        "PostgreSQL assure la persistance. Les fichiers uploadés (ex. photos de services) sont servis "
        "comme ressources statiques sous le chemin /uploads.",
    )
    doc.add_paragraph(
        "Schéma logique (vue d’ensemble) :\n\n"
        "[ Application Flutter ]  --HTTP/JSON-->  [ API Express :3000 ]\n"
        "                                              |\n"
        "                                              v\n"
        "                                        [ PostgreSQL ]",
        style="No Spacing",
    )

    add_heading(doc, "3.2 Choix technologiques", 2)
    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Couche"
    hdr[1].text = "Technologie"
    hdr[2].text = "Justification"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("Mobile", "Flutter (Dart 3.x)", "UI multiplateforme, écosystème mature, intégration HTTP et state management (Provider)."),
        ("API", "Node.js + Express 4", "Légèreté, déploiement simple, écosystème npm pour JWT, multer, pg."),
        ("Données", "PostgreSQL + pg", "Modèle relationnel, contraintes d’intégrité, requêtes SQL pour agrégations réservations / avis."),
        ("Auth", "JWT + bcryptjs", "Stateless pour API REST ; hachage des mots de passe."),
    ]
    for a, b, c in rows:
        row = tbl.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    add_heading(doc, "3.3 Modèle de données (résumé)", 2)
    add_para(
        doc,
        "Les entités principales sont : utilisateurs (users) avec rôle client, prestataire ou admin ; "
        "catégories de métiers (categories) ; services proposés par prestataire (services) ; "
        "réservations (reservations) liant client, prestataire et service ; avis (avis) liés à une réservation. "
        "Les statuts de réservation incluent : en_attente, confirme, en_cours, termine, annule. "
        "Les statuts de paiement prévus : non_paye, paye, rembourse.",
    )

    # --- 4. Réalisation ---
    add_heading(doc, "4. Réalisation", 1)
    add_heading(doc, "4.1 Organisation du dépôt", 2)
    add_para(
        doc,
        "Le code est organisé sous servidom-app/ avec au minimum : dossier backend/ (API Node, fichier server.js, "
        "dossiers routes, controllers, middleware, config) et dossier mobile/ (projet Flutter lib/, écrans, providers).",
    )
    add_heading(doc, "4.2 Modules backend", 2)
    add_bullets(
        doc,
        [
            "auth : inscription, connexion, endpoint profil courant (/me) avec JWT.",
            "users : mise à jour du profil authentifié (PUT /api/users/profil).",
            "services : catégories, liste prestataires, détail prestataire, création de service avec upload.",
            "reservations : création, liste, mise à jour de statut, avis.",
        ],
    )
    add_heading(doc, "4.3 Sécurité", 2)
    add_para(
        doc,
        "Les mots de passe ne sont jamais stockés en clair. Le middleware d’authentification vérifie le jeton "
        "sur les routes sensibles. L’inscription publique n’accepte que les rôles client et prestataire.",
    )

    # --- 5. Déploiement ---
    add_heading(doc, "5. Déploiement et exploitation", 1)
    add_heading(doc, "5.1 Prérequis", 2)
    add_bullets(
        doc,
        [
            "Node.js (LTS recommandé) et npm.",
            "PostgreSQL installé et base de données créée.",
            "Flutter SDK pour compiler et exécuter l’application mobile.",
        ],
    )
    add_heading(doc, "5.2 Variables d’environnement (backend)", 2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "Variable"
    t2.rows[0].cells[1].text = "Rôle"
    for c in t2.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    env_rows = [
        ("DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD", "Connexion PostgreSQL"),
        ("JWT_SECRET, JWT_EXPIRES_IN", "Signature et durée de vie du jeton JWT"),
        ("PORT", "Port d’écoute de l’API (défaut 3000 si absent)"),
    ]
    for k, v in env_rows:
        rr = t2.add_row().cells
        rr[0].text = k
        rr[1].text = v

    add_heading(doc, "5.3 Initialisation de la base", 2)
    add_para(
        doc,
        "Après configuration du fichier .env dans backend/, installer les dépendances (npm install) puis exécuter "
        "le script Node src/config/initDB.js depuis le répertoire backend : cela crée les tables et insère les "
        "catégories de référence (Cuisinière, Ménagère, Jardinier, Plombier, Électricien, Sécurité).",
    )
    add_heading(doc, "5.4 Démarrage", 2)
    add_bullets(
        doc,
        [
            "Backend : npm run dev ou npm start (répertoire backend/).",
            "Mobile : flutter pub get puis flutter run en pointant l’URL de l’API selon la configuration du projet.",
        ],
    )

    # --- 6. Tests ---
    add_heading(doc, "6. Tests et validation", 1)
    add_para(
        doc,
        "Une stratégie de validation manuelle recommandée : parcours inscription client et prestataire ; "
        "connexion ; création d’un service avec image ; création de réservation côté client ; "
        "consultation des réservations des deux côtés ; changement de statut ; soumission d’un avis. "
        "Des tests automatisés (unitaires / intégration) peuvent être ajoutés sur le backend et les widgets Flutter.",
    )

    # --- 7. Difficultés ---
    add_heading(doc, "7. Difficultés rencontrées et solutions", 1)
    add_para(
        doc,
        "[À compléter par l’équipe : décrire par exemple la configuration réseau émulateur / appareil physique, "
        "les problèmes CORS ou de certificats, la migration du schéma SQL, etc.]",
    )

    # --- 8. Conclusion ---
    add_heading(doc, "8. Conclusion et perspectives", 1)
    add_para(
        doc,
        "ServiDom fournit un socle fonctionnel pour la digitalisation des services à domicile. "
        "Les évolutions possibles incluent : tableau de bord administrateur ; notifications push ; "
        "intégration d’un prestataire de paiement ; géolocalisation avancée ; système de messagerie in-app ; "
        "renforcement des tests et monitoring en production.",
    )

    # --- Annexes ---
    doc.add_page_break()
    add_heading(doc, "Annexes", 1)
    add_heading(doc, "Annexe A – Principales routes API", 2)
    api = doc.add_table(rows=1, cols=4)
    h = api.rows[0].cells
    h[0].text = "Méthode"
    h[1].text = "Chemin"
    h[2].text = "Auth"
    h[3].text = "Description"
    for c in h:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    routes = [
        ("POST", "/api/auth/register", "Non", "Inscription client ou prestataire"),
        ("POST", "/api/auth/login", "Non", "Connexion, retour JWT"),
        ("GET", "/api/auth/me", "Oui", "Utilisateur courant"),
        ("PUT", "/api/users/profil", "Oui", "Mise à jour profil"),
        ("GET", "/api/services/categories", "Non", "Liste des catégories"),
        ("GET", "/api/services/prestataires", "Non", "Liste des prestataires"),
        ("GET", "/api/services/prestataires/:id", "Non", "Détail d’un prestataire"),
        ("POST", "/api/services/", "Oui", "Création d’un service (+ upload image)"),
        ("POST", "/api/reservations/", "Oui (client)", "Création réservation"),
        ("GET", "/api/reservations/mes-reservations", "Oui", "Liste selon le rôle"),
        ("PATCH", "/api/reservations/:id/statut", "Oui", "Mise à jour du statut"),
        ("POST", "/api/reservations/avis", "Oui", "Déposer un avis"),
        ("GET", "/", "Non", "Contrôle santé API"),
    ]
    for row_data in routes:
        row = api.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    add_heading(doc, "Annexe B – Figures", 2)
    add_para(
        doc,
        "Emplacements réservés pour captures d’écran de l’application et diagrammes UML détaillés "
        "(ex. fichier plantuml_export.puml du dépôt).",
    )

    add_heading(doc, "Glossaire", 2)
    gloss = [
        ("API", "Interface de programmation ; ici REST sur HTTP avec corps JSON."),
        ("JWT", "JSON Web Token : jeton signé pour authentifier les requêtes."),
        ("ORM", "Object-Relational Mapping ; le projet utilise SQL direct via le client pg."),
    ]
    for term, defn in gloss:
        p = doc.add_paragraph()
        p.add_run(term + " : ").bold = True
        p.add_run(defn)

    add_heading(doc, "Bibliographie et liens", 2)
    add_bullets(
        doc,
        [
            "Documentation Flutter : https://docs.flutter.dev/",
            "Express : https://expressjs.com/",
            "PostgreSQL : https://www.postgresql.org/docs/",
            "Node.js : https://nodejs.org/docs/",
        ],
    )

    doc.save(out_path)
    print(f"Document généré : {out_path}")


if __name__ == "__main__":
    main()
